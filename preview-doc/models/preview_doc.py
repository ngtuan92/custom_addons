import base64
import io
import os
import subprocess
import tempfile
import logging

from odoo import models, fields, api, _
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)


class MyModel(models.Model):
    _name = "my.model"
    _description = "Word Preview Export Demo"

    name = fields.Char(string="Name", required=True)

    word_file = fields.Binary(
        string="Word file (.docx)",
        attachment=True
    )
    word_filename = fields.Char(string="Word filename")

    pdf_preview = fields.Binary(
        string="PDF preview",
        attachment=True,
        readonly=True
    )
    pdf_filename = fields.Char(
        string="PDF filename",
        readonly=True
    )

    def write(self, vals):
        """Override write để auto convert sau khi Save"""
        result = super().write(vals)

        # Convert sau khi Save nếu có word_file mới
        if 'word_file' in vals or 'word_filename' in vals:
            for record in self:
                if record.word_file and record.word_filename and record.word_filename.lower().endswith('.docx'):
                    try:
                        _logger.info(f"=== AUTO CONVERT ON SAVE: {record.word_filename} ===")
                        record._do_convert()
                    except Exception as e:
                        _logger.error(f"Auto convert failed: {str(e)}", exc_info=True)

        return result

    @api.model_create_multi
    def create(self, vals_list):
        """Override create để auto convert khi tạo mới"""
        records = super().create(vals_list)

        for record in records:
            if record.word_file and record.word_filename and record.word_filename.lower().endswith('.docx'):
                try:
                    _logger.info(f"=== AUTO CONVERT ON CREATE: {record.word_filename} ===")
                    record._do_convert()
                except Exception as e:
                    _logger.error(f"Auto convert failed: {str(e)}", exc_info=True)

        return records

    def _do_convert(self):
        """Internal method to convert - KHÔNG dùng @api.onchange"""
        self.ensure_one()

        _logger.info(f"Converting: {self.word_filename}")

        # Decode
        docx_bytes = base64.b64decode(self.word_file)
        _logger.info(f"Decoded: {len(docx_bytes)} bytes")

        # Convert
        pdf_bytes = self._convert_docx_to_pdf(docx_bytes)
        _logger.info(f"PDF generated: {len(pdf_bytes)} bytes")

        # Update fields - SỬ DỤNG SUDO để bypass readonly
        base_name = self.word_filename.rsplit(".", 1)[0]
        self.sudo().write({
            'pdf_preview': base64.b64encode(pdf_bytes),
            'pdf_filename': f"{base_name}.pdf"
        })

        _logger.info(f"=== SAVED TO DB: {base_name}.pdf ===")

    def _convert_docx_to_pdf(self, docx_bytes):
        """Convert DOCX to PDF using LibreOffice"""
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "input.docx")
            pdf_path = os.path.join(tmpdir, "input.pdf")

            with open(docx_path, "wb") as f:
                f.write(docx_bytes)

            soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

            if not os.path.exists(soffice_path):
                raise UserError(f"LibreOffice not found: {soffice_path}")

            cmd = [
                soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", tmpdir,
                docx_path
            ]

            subprocess.run(cmd, check=True, stdout=subprocess.PIPE,
                           stderr=subprocess.PIPE, timeout=60)

            if not os.path.exists(pdf_path):
                raise UserError("PDF conversion failed")

            with open(pdf_path, "rb") as f:
                return f.read()

    def action_export_word(self):
        """Export Word document"""
        self.ensure_one()

        try:
            from docx import Document
        except:
            raise UserError("Cài python-docx: pip install python-docx")

        doc = Document()
        doc.add_heading("DEMO EXPORT WORD", level=1)
        doc.add_paragraph(f"Tên: {self.name}")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        attachment = self.env["ir.attachment"].create({
            "name": f"{self.name}.docx",
            "type": "binary",
            "datas": base64.b64encode(buffer.read()),
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "res_model": self._name,
            "res_id": self.id,
        })

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }
